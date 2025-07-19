import os
import re
import sys
import subprocess
import tkinter as tk

from pathlib import Path
from tkinter import filedialog, messagebox, ttk
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.oxml.ns import qn

IMAGE_EXTENSIONS = ('.jpg', '.arw', '.jpeg', '.png', '.tif', '.tiff', '.cr2', '.nef', '.dng')


def natural_sort_key(s):
    return [int(text) if text.isdigit() else text.lower() for text in re.split(r'(\d+)', str(s))]


def fill_metadata(table, metadata, num, font_name="黑体", font_size=10.5):
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if cell.text.strip() in metadata and idx + num < len(row.cells):
                # 清空目标单元格，重新写入带字体样式的内容
                cell_to_fill = row.cells[idx + num]
                cell_to_fill.text = ""
                p = cell_to_fill.paragraphs[0]
                run = p.add_run(metadata[cell.text.strip()])
                run.font.name = font_name
                run.font.size = Pt(font_size)
                r = run._element.rPr.rFonts
                r.set(qn('w:eastAsia'), font_name)


def center_table_text(document):
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def process_folders_and_update_word(input_base_dir, template_path, output_path,
                                    processor_name, filler_name, date_value,
                                    work_location, unit_name, progress_var, progress_label, progress_bar,
                                    font_name="黑体", font_size=10.5):
    def set_cell_text(cell, text, font_name=font_name, font_size=font_size):
        cell.text = ""  # 清空原内容
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.name = font_name
        run.font.size = Pt(font_size)
        r = run._element.rPr.rFonts
        r.set(qn('w:eastAsia'), font_name)  # 设置中文字体兼容

    folder_data = []
    folders = sorted([f for f in Path(input_base_dir).iterdir() if f.is_dir()],
                     key=natural_sort_key)

    total_folders = len(folders)
    if total_folders == 0:
        messagebox.showerror("错误", "输入目录中没有子文件夹。")
        return

    for idx, folder_path in enumerate(folders, start=1):
        image_files = [item.name for item in folder_path.iterdir()
                       if item.is_file() and item.suffix.lower() in IMAGE_EXTENSIONS]
        if image_files:
            sorted_image_files = sorted(image_files, key=natural_sort_key)
            first_file = Path(sorted_image_files[0]).stem
            second_file = Path(sorted_image_files[1]).stem if len(sorted_image_files) >= 2 else ""
            last_file = Path(sorted_image_files[-1]).stem if len(sorted_image_files) >= 2 else ""
            second_to_last_combined = f"{second_file}-{last_file}" if second_file and last_file else ""
            file_count_second_to_last = max(0, len(sorted_image_files) - 1)
            total_folder_size_bytes = sum(
                item.stat().st_size for item in folder_path.iterdir()
                if item.is_file() and item.suffix.lower() in IMAGE_EXTENSIONS
            )
            total_folder_size_gb = total_folder_size_bytes / (1024 ** 3)

            folder_data.append({
                "first_file": first_file,
                "second_to_last_combined": second_to_last_combined,
                "file_count": file_count_second_to_last,
                "folder_size_gb": total_folder_size_gb,
                "processor_name": processor_name
            })

        percent = int((idx / total_folders) * 100)
        progress_var.set(percent)
        progress_label.config(text=f"{percent}%")
        progress_bar.update()

    if not folder_data:
        messagebox.showwarning("提示", "没有收集到任何文件夹数据。")
        return

    document = Document(template_path)
    if not document.tables:
        messagebox.showerror("错误", "模板中没有表格。")
        return

    table = document.tables[0]
    fill_metadata(table, {"填表人": filler_name}, 1)
    fill_metadata(table, {"日期": date_value, "工作地点": work_location}, 2)
    fill_metadata(table, {"文保单位名称及标号": unit_name}, 3)

    start_row_idx = 7
    if len(table.rows) <= start_row_idx:
        for _ in range(start_row_idx + 1 - len(table.rows)):
            table.add_row()

    for i, data_entry in enumerate(folder_data):
        row_idx = start_row_idx + i
        row_cells = table.rows[row_idx].cells if row_idx < len(table.rows) else table.add_row().cells
        try:
            if len(row_cells) >= 2:
                row_cells[0].merge(row_cells[1])
            if len(row_cells) >= 5:
                row_cells[2].merge(row_cells[4])
            if len(row_cells) >= 7:
                row_cells[5].merge(row_cells[6])
            if len(row_cells) >= 9:
                row_cells[7].merge(row_cells[8])
            if len(row_cells) >= 11:
                row_cells[9].merge(row_cells[10])
        except:
            pass

        set_cell_text(row_cells[0], data_entry["first_file"])
        set_cell_text(row_cells[2], data_entry["second_to_last_combined"])
        set_cell_text(row_cells[5], str(data_entry["file_count"]))
        set_cell_text(row_cells[7], f"{data_entry['folder_size_gb']:.2f} GB")
        set_cell_text(row_cells[9], data_entry["processor_name"])

    center_table_text(document)
    document.save(output_path)

    folder_to_open = os.path.dirname(output_path)
    if os.name == 'nt':
        os.startfile(folder_to_open)
    else:
        subprocess.Popen(['open', folder_to_open])

    messagebox.showinfo("完成", f"报告已生成：{output_path}")


def run_app():
    def browse_template():
        path = filedialog.askopenfilename(title="选择模板文件", filetypes=[("Word 文件", "*.docx")])
        entry_template.delete(0, tk.END)
        entry_template.insert(0, path)

    def browse_folder():
        path = filedialog.askdirectory(title="选择目标图片文件夹")
        entry_folder.delete(0, tk.END)
        entry_folder.insert(0, path)

    def generate_report():
        template = entry_template.get().strip()
        folder = entry_folder.get().strip()
        output_name = entry_output.get().strip()
        filler = entry_filler.get().strip()
        date_val = entry_date.get().strip()
        location = entry_location.get().strip()
        unit = entry_unit.get().strip()
        processor = entry_processor.get().strip()

        if not all([template, folder, output_name, filler, date_val, location, unit, processor]):
            messagebox.showwarning("提示", "请填写所有字段！")
            return

        # output_path = os.path.join(os.path.dirname(folder), f"{output_name}.docx")
        exe_dir = os.path.dirname(sys.executable)  # 打包exe时用
        # exe_dir = os.path.dirname(os.path.abspath(__file__))  # 调试运行时用

        output_path = os.path.join(exe_dir, f"{output_name}.docx")
        progress_var.set(0)
        process_folders_and_update_word(folder, template, output_path,
                                        processor, filler, date_val, location, unit,
                                        progress_var, progress_label, progress_bar)
        root.destroy()

    root = tk.Tk()
    root.title("Word 报告生成器")
    window_width, window_height = 650, 500
    screen_width = root.winfo_screenwidth()
    screen_height = root.winfo_screenheight()
    x = int((screen_width / 2) - (window_width / 2))
    y = int((screen_height / 2) - (window_height / 2))
    root.geometry(f"{window_width}x{window_height}+{x}+{y}")

    def add_row(label_text, entry_var, browse_cmd=None):
        frame = tk.Frame(root)
        frame.pack(fill="x", pady=5)
        tk.Label(frame, text=label_text, width=22, anchor="w").pack(side="left")
        entry = tk.Entry(frame, textvariable=entry_var, width=40)
        entry.pack(side="left", expand=True, fill="x")
        if browse_cmd:
            tk.Button(frame, text="浏览", command=browse_cmd).pack(side="right")
        return entry

    entry_template_var = tk.StringVar()
    entry_folder_var = tk.StringVar()
    entry_output_var = tk.StringVar()
    entry_filler_var = tk.StringVar()
    entry_date_var = tk.StringVar()
    entry_location_var = tk.StringVar()
    entry_unit_var = tk.StringVar()
    entry_processor_var = tk.StringVar()

    entry_template = add_row("模板文件：", entry_template_var, browse_template)
    entry_folder = add_row("图片文件夹：", entry_folder_var, browse_folder)
    entry_output = add_row("输出文件名：", entry_output_var)
    entry_filler = add_row("填表人：", entry_filler_var)
    entry_date = add_row("日期 (YYYY-MM-DD)：", entry_date_var)
    entry_location = add_row("工作地点：", entry_location_var)
    entry_unit = add_row("文保单位名称及标号：", entry_unit_var)
    entry_processor = add_row("校色人员：", entry_processor_var)

    progress_frame = tk.Frame(root)
    progress_frame.pack(fill="x", padx=20, pady=10)
    progress_var = tk.IntVar()
    progress_bar = ttk.Progressbar(progress_frame, variable=progress_var, maximum=100)
    progress_bar.pack(side="left", fill="x", expand=True)
    progress_label = tk.Label(progress_frame, text="0%")
    progress_label.pack(side="right")

    tk.Button(root, text="生成报告", command=generate_report, bg="#4CAF50", fg="white").pack(pady=10)

    root.mainloop()


if __name__ == "__main__":
    run_app()
